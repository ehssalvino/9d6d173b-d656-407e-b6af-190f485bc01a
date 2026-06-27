from __future__ import annotations

import re
import shutil
import tempfile
import urllib.request
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from lxml import etree

from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_P = f"{{{W_NS}}}p"
W_T = f"{{{W_NS}}}t"
TAG_RE = re.compile(r"\{\{[^{}]+\}\}")
MONTHS = ["JAN", "FEV", "MAR", "ABR", "MAI", "JUN", "JUL", "AGO", "SET", "OUT", "NOV", "DEZ"]
# Valores da aba Dimensionamento, células R34:R45, da planilha de referência.
MONTHLY_IRRADIATION = [5.58, 5.85, 4.86, 4.42, 3.82, 3.67, 3.91, 4.85, 5.06, 5.21, 4.93, 5.50]


class GenerationError(RuntimeError):
    pass


@dataclass
class GenerationResult:
    output: Path
    replaced_tags: int
    missing_values: list[str]
    unresolved_tags: list[str]


def _number(value: Any) -> float | None:
    if value in (None, ""):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace(" ", "")
    if "," in text and "." in text:
        text = text.replace(".", "").replace(",", ".")
    else:
        text = text.replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return None


def _fmt(value: Any, placeholder: str = "") -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, bool):
        return "Sim" if value else "Não"
    if isinstance(value, float):
        if "PERC" in placeholder:
            return f"{value * 100:.2f}".replace(".", ",")
        return f"{value:.6f}".rstrip("0").rstrip(".").replace(".", ",")
    return str(value)


def _safe_div(a: float | None, b: float | None) -> float | None:
    if a is None or b in (None, 0):
        return None
    return a / b


def _default_loads(target_kwp: float | None) -> list[tuple[str, float, float, float]]:
    # Residência média. A soma fica próxima de 5,7 kW, adequada para memoriais
    # residenciais pequenos; o usuário pode sobrescrever qualquer linha no Excel.
    base = [
        ("Iluminação LED", 10, 12, 0.92),
        ("Refrigerador", 300, 1, 0.85),
        ("Televisores", 100, 2, 0.90),
        ("Ventiladores", 100, 2, 0.90),
        ("Máquina de lavar", 500, 1, 0.85),
        ("Micro-ondas", 1200, 1, 0.92),
        ("Computador/notebook", 200, 1, 0.90),
        ("Ar-condicionado", 1200, 1, 0.90),
        ("Bomba d'água", 370, 1, 0.82),
        ("Tomadas de uso geral", 600, 1, 0.90),
        ("Roteador/eletrônicos", 30, 1, 0.90),
        ("Ferro elétrico", 1000, 1, 0.95),
    ]
    if not target_kwp or target_kwp <= 0:
        return base
    target_w = max(3500, min(target_kwp * 1000, 8000))
    total_w = sum(power * qty for _, power, qty, _ in base)
    if total_w <= target_w:
        return base
    scale = target_w / total_w
    scaled = []
    for desc, power, qty, fp in base:
        if desc in {"Refrigerador", "Roteador/eletrônicos"}:
            scaled.append((desc, power, qty, fp))
        else:
            scaled.append((desc, max(10, round(power * scale / 10) * 10), qty, fp))
    return scaled


def _word_number(value: int | float | None) -> str:
    words = {
        1: "UM",
        2: "DOIS",
        3: "TRÊS",
        4: "QUATRO",
        5: "CINCO",
        6: "SEIS",
        7: "SETE",
        8: "OITO",
        9: "NOVE",
        10: "DEZ",
        11: "ONZE",
        12: "DOZE",
    }
    if value is None:
        return ""
    return words.get(int(value), str(int(value)))


def _connection_defaults(inp: dict[str, Any]) -> dict[str, Any]:
    raw = str(inp.get("TIPO_LIGACAO") or inp.get("INV_CONEXAO") or "").upper()
    if "TRI" in raw or "3" in raw:
        phases, kind = 3, "TRIFÁSICO"
    elif "BI" in raw or "2" in raw:
        phases, kind = 2, "BIFÁSICO"
    else:
        phases, kind = 1, "MONOFÁSICO"
    return {
        "{{TIPO_LIGACAO}}": kind,
        "{{CONDUTORES_FASE_QTD}}": phases,
        "{{CONDUTORES_FASE_QTD_TEXTO}}": _word_number(phases),
        "{{CONDUTORES_TOTAL_TEXTO}}": _word_number(phases + 1),
        "{{CONDUTOR_FASE_MM2}}": inp.get("CONDUTOR_FASE_MM2") or 10,
        "{{CONDUTOR_NEUTRO_DESC}}": inp.get("CONDUTOR_NEUTRO_DESC") or "condutor NEUTRO",
        "{{CONCESSIONARIA}}": inp.get("CONCESSIONARIA") or "LIGHT S.A.",
        "{{ESTADO}}": inp.get("ESTADO") or "RIO DE JANEIRO",
        "{{DISJ_POLOS}}": inp.get("DISJ_POLOS") or phases,
        "{{DISJ_CORRENTE_A}}": inp.get("DISJ_CORRENTE_A") or 40,
        "{{DISJ_TENSAO_V}}": inp.get("DISJ_TENSAO_V") or inp.get("INV_VAC_NOM") or 220,
    }


def _calculated_values(workbook) -> dict[str, Any]:
    inp = {
        workbook["Inputs"][f"A{row}"].value: workbook["Inputs"][f"B{row}"].value
        for row in range(2, workbook["Inputs"].max_row + 1)
    }
    n = lambda key: _number(inp.get(key))

    consumption = [_number(workbook["Consumo_24m"][f"B{row}"].value) for row in range(2, 26)]
    valid_consumption = [value for value in consumption if value is not None]

    module_qty = n("MOD_QTD") or 0
    module_wp = n("MOD_WP")
    target_kwp = module_qty * module_wp / 1000 if module_qty and module_wp else n("FV_POT_KWP")
    default_loads = _default_loads(target_kwp)

    load_desc: list[Any] = []
    load_power: list[Any] = []
    load_quantity: list[Any] = []
    load_fp: list[Any] = []
    load_kw: list[float | None] = []
    load_kva: list[float | None] = []
    for row in range(2, 26):
        default = default_loads[row - 2] if row - 2 < len(default_loads) else None
        desc = workbook["Cargas_24"][f"B{row}"].value
        power = _number(workbook["Cargas_24"][f"C{row}"].value)
        quantity = _number(workbook["Cargas_24"][f"D{row}"].value)
        power_factor = _number(workbook["Cargas_24"][f"E{row}"].value)
        if default and desc in (None, "") and power is None and quantity is None:
            desc, power, quantity, power_factor = default
        load_desc.append(desc)
        load_power.append(power)
        load_quantity.append(quantity)
        load_fp.append(power_factor)
        kw = power * quantity / 1000 if power is not None and quantity is not None else None
        load_kw.append(kw)
        load_kva.append(_safe_div(kw, power_factor))

    values: dict[str, Any] = {}
    for index, value in enumerate(consumption, 1):
        values[f"{{{{CONS_MES{index}}}}}"] = value
    values["{{CONS_TOTAL}}"] = sum(valid_consumption) if valid_consumption else None
    values["{{CONS_MEDIA}}"] = (
        sum(valid_consumption) / len(valid_consumption) if valid_consumption else None
    )

    for index, (kw, kva) in enumerate(zip(load_kw, load_kva), 1):
        values[f"{{{{CARGA{index}_DESC}}}}"] = load_desc[index - 1]
        values[f"{{{{CARGA{index}_P_W}}}}"] = load_power[index - 1]
        values[f"{{{{CARGA{index}_QTD}}}}"] = load_quantity[index - 1]
        values[f"{{{{CARGA{index}_FP}}}}"] = load_fp[index - 1]
        values[f"{{{{CARGA{index}_CI_KW}}}}"] = kw
        values[f"{{{{CARGA{index}_CI_KVA}}}}"] = kva
    values["{{CARGA_TOTAL_KW}}"] = sum(value for value in load_kw if value is not None)
    values["{{CARGA_TOTAL_KVA}}"] = sum(value for value in load_kva if value is not None)

    design_strings = 2 if module_qty >= 10 else 1
    modules_per_string = int((module_qty + design_strings - 1) // design_strings) if module_qty else module_qty
    if n("INV_STRINGS") is not None:
        values["{{INV_STRINGS}}"] = n("INV_STRINGS")
    values["{{STRING_QTD}}"] = design_strings
    values["{{MOD_QTD_TEXTO}}"] = _word_number(module_qty)
    values["{{INV_QTD_TEXTO}}"] = _word_number(n("INV_QTD"))
    inv_qty = n("INV_QTD")
    if inv_qty == 1:
        values["{{INV_DESC}}"] = "01 (um) inversor"
    elif inv_qty:
        values["{{INV_DESC}}"] = f"{int(inv_qty):02d} ({_word_number(inv_qty).lower()}) inversores"
    else:
        values["{{INV_DESC}}"] = "inversor(es)"
    values["{{FV_POT_STC_KWP}}"] = target_kwp
    values["{{FV_POT_VI_W}}"] = (
        n("MOD_VMP") * n("MOD_IMP") * module_qty
        if all(n(key) is not None for key in ("MOD_VMP", "MOD_IMP")) and module_qty
        else None
    )
    values["{{FV_POT_VI_KW}}"] = values["{{FV_POT_VI_W}}"] / 1000 if values["{{FV_POT_VI_W}}"] else None
    values["{{VCC_NOM}}"] = (
        n("MOD_VMP") * modules_per_string if n("MOD_VMP") is not None and modules_per_string else None
    )
    values["{{MODULOS_POR_STRING}}"] = modules_per_string
    values["{{ICC_MAX}}"] = n("MOD_IMP")
    values["{{IAC_MAX}}"] = _safe_div(
        n("INV_PCA_KW") * 1000 if n("INV_PCA_KW") is not None else None,
        n("INV_VAC_NOM"),
    )
    values["{{DV_CC_V}}"] = (
        n("MOD_IMP") * n("R_CC_OHMKM") * (2 * n("DIST_CC_M") / 1000)
        if all(n(key) is not None for key in ("MOD_IMP", "R_CC_OHMKM", "DIST_CC_M"))
        else None
    )
    values["{{DIST_CC_KM}}"] = n("DIST_CC_M") / 1000 if n("DIST_CC_M") is not None else None
    values["{{DV_CC_PERC}}"] = _safe_div(values["{{DV_CC_V}}"], values["{{VCC_NOM}}"])
    values["{{DV_CA_V}}"] = (
        values["{{IAC_MAX}}"] * n("R_CA_OHMKM") * (2 * n("DIST_CA_M") / 1000)
        if values["{{IAC_MAX}}"] is not None
        and n("R_CA_OHMKM") is not None
        and n("DIST_CA_M") is not None
        else None
    )
    values["{{DIST_CA_KM}}"] = n("DIST_CA_M") / 1000 if n("DIST_CA_M") is not None else None
    values["{{DV_CA_PERC}}"] = _safe_div(values["{{DV_CA_V}}"], n("INV_VAC_NOM"))
    values.update(_connection_defaults(inp))
    optional_defaults = {
        "{{MOD_TIPO_CELULA}}": inp.get("MOD_TIPO_CELULA") or "",
        "{{MOD_N_CELULAS}}": inp.get("MOD_N_CELULAS") or "",
        "{{MOD_DIMENSOES}}": inp.get("MOD_DIMENSOES") or "",
        "{{MOD_PESO_KG}}": inp.get("MOD_PESO_KG") or "",
        "{{INV_PCA_MAX_KW}}": inp.get("INV_PCA_MAX_KW") or inp.get("INV_PCA_KW") or "",
        "{{INV_VCA_MAX}}": inp.get("INV_VCA_MAX") or "—",
        "{{INV_VCA_MIN}}": inp.get("INV_VCA_MIN") or "—",
        "{{INV_THD}}": inp.get("INV_THD") or "≤ 3",
        "{{PROT_DISJ_POLOS}}": inp.get("PROT_DISJ_POLOS") or 4,
        "{{PROT_DISJ_TENSAO_V}}": inp.get("PROT_DISJ_TENSAO_V") or 1000,
        "{{PROT_DISJ_CORRENTE_A}}": inp.get("PROT_DISJ_CORRENTE_A") or 16,
        "{{PROT_DISJ_FREQ_HZ}}": inp.get("PROT_DISJ_FREQ_HZ") or 60,
        "{{PROT_DISJ_CAP_INT_KA}}": inp.get("PROT_DISJ_CAP_INT_KA") or 40,
        "{{PROT_DISJ_CURVA}}": inp.get("PROT_DISJ_CURVA") or "Classe 2",
        "{{DSV_POLOS}}": inp.get("DSV_POLOS") or 4,
        "{{DSV_TENSAO_V}}": inp.get("DSV_TENSAO_V") or 1000,
        "{{DSV_CORRENTE_A}}": inp.get("DSV_CORRENTE_A") or 32,
        "{{DSV_FECHAMENTO_CC_A}}": inp.get("DSV_FECHAMENTO_CC_A") or 600,
        "{{DSV_CAP_INT_KA}}": inp.get("DSV_CAP_INT_KA") or 40,
        "{{DSV_CURVA}}": inp.get("DSV_CURVA") or "Classe 2",
        "{{ATERRAMENTO_ESQUEMA}}": inp.get("ATERRAMENTO_ESQUEMA") or "TT",
        "{{ATERRAMENTO_RESISTENCIA_OHM}}": inp.get("ATERRAMENTO_RESISTENCIA_OHM") or 15,
        "{{ATERRAMENTO_HASTES_QTD}}": inp.get("ATERRAMENTO_HASTES_QTD") or 1,
        "{{ATERRAMENTO_HASTE_M}}": inp.get("ATERRAMENTO_HASTE_M") or "2,44",
        "{{ATERRAMENTO_HASTE_SECAO}}": inp.get("ATERRAMENTO_HASTE_SECAO") or "5/8",
        "{{ATERRAMENTO_BITOLA_MM2}}": inp.get("ATERRAMENTO_BITOLA_MM2") or 6,
        "{{CABO_ISOLACAO}}": inp.get("CABO_ISOLACAO") or "XLPE",
        "{{CABO_ISOLAMENTO_KV}}": inp.get("CABO_ISOLAMENTO_KV") or "0,6/1",
        "{{CABO_BITOLA_MM2}}": inp.get("CABO_BITOLA_MM2") or inp.get("BITOLA_CA_MM2") or 6,
        "{{CABO_CAPACIDADE_A}}": inp.get("CABO_CAPACIDADE_A") or 53,
        "{{CABO_METODO_INSTALACAO}}": inp.get("CABO_METODO_INSTALACAO") or "B1 (eletroduto aparente e cabos unipolares)",
        "{{FATOR_TEMPERATURA}}": inp.get("FATOR_TEMPERATURA") or "0,71",
        "{{FATOR_AGRUPAMENTO}}": inp.get("FATOR_AGRUPAMENTO") or "0,57",
        "{{TEMPERATURA_AMBIENTE_C}}": inp.get("TEMPERATURA_AMBIENTE_C") or 50,
    }
    values.update(optional_defaults)
    return values


def _read_values(excel_path: Path) -> tuple[dict[str, str], list[str]]:
    formulas = load_workbook(excel_path, data_only=False)
    cached = load_workbook(excel_path, data_only=True)
    if "Mapeamento" not in formulas.sheetnames:
        raise GenerationError("A planilha não possui a aba 'Mapeamento'.")

    calculated = _calculated_values(formulas)
    replacements: dict[str, str] = {}
    missing: list[str] = []

    for placeholder, origin in formulas["Mapeamento"].iter_rows(min_row=2, max_col=2, values_only=True):
        if not placeholder or not origin or placeholder in {"{{FIG_LOCALIZACAO}}", "{{GRAFICO_GERACAO}}"}:
            continue
        try:
            sheet, cell = str(origin).split("!", 1)
            value = cached[sheet][cell].value
            if value in (None, "", "#DIV/0!") and placeholder in calculated:
                value = calculated[placeholder]
            if value in (None, "", "#DIV/0!"):
                value = formulas[sheet][cell].value
                if isinstance(value, str) and value.startswith("="):
                    value = None
        except (KeyError, ValueError) as exc:
            raise GenerationError(f"Origem inválida no mapeamento: {origin}") from exc
        replacements[str(placeholder).strip()] = _fmt(value, str(placeholder))
        if value in (None, ""):
            missing.append(str(placeholder))
    for placeholder, value in calculated.items():
        replacements.setdefault(placeholder, _fmt(value, placeholder))
    return replacements, missing


def _replace_in_paragraph(paragraph: etree._Element, replacements: dict[str, str]) -> int:
    nodes = list(paragraph.iter(W_T))
    if not nodes:
        return 0
    text = "".join(node.text or "" for node in nodes)
    if "{{" not in text:
        return 0

    count = 0
    for placeholder, value in replacements.items():
        while placeholder in text:
            start = text.index(placeholder)
            end = start + len(placeholder)
            offsets: list[tuple[int, int, ET.Element]] = []
            cursor = 0
            for node in nodes:
                node_text = node.text or ""
                offsets.append((cursor, cursor + len(node_text), node))
                cursor += len(node_text)
            first = True
            for node_start, node_end, node in offsets:
                if node_end <= start or node_start >= end:
                    continue
                local_start = max(0, start - node_start)
                local_end = min(node_end, end) - node_start
                node_text = node.text or ""
                insertion = value if first else ""
                node.text = node_text[:local_start] + insertion + node_text[local_end:]
                first = False
            text = "".join(node.text or "" for node in nodes)
            count += 1
    return count


def _replace_all_xml_parts(docx_path: Path, replacements: dict[str, str]) -> int:
    temp_path = docx_path.with_suffix(".tmp.docx")
    replaced = 0
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                    for paragraph in root.iter(W_P):
                        replaced += _replace_in_paragraph(paragraph, replacements)
                    data = etree.tostring(root, encoding="utf-8", xml_declaration=True, standalone=True)
                except etree.XMLSyntaxError:
                    pass
            target.writestr(item, data)
    temp_path.replace(docx_path)
    return replaced


def _make_location_figure(latitude: Any, longitude: Any, output: Path) -> Path:
    lat_num = _number(latitude)
    lon_num = _number(longitude)
    if lat_num is not None and lon_num is not None:
        lat_map = -abs(lat_num) if lat_num > 0 else lat_num
        lon_map = -abs(lon_num) if lon_num > 0 else lon_num
        url = (
            "https://staticmap.openstreetmap.de/staticmap.php"
            f"?center={lat_map},{lon_map}&zoom=18&size=900x520&markers={lat_map},{lon_map},red-pushpin"
        )
        try:
            with urllib.request.urlopen(url, timeout=12) as response:
                image = Image.open(response).convert("RGB")
                image.save(output)
                return output
        except Exception:
            pass

    width, height = 1200, 620
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    for x in range(50, 760, 50):
        draw.line([(x, 50), (x, 570)], fill=(222, 226, 230), width=1)
    for y in range(50, 600, 50):
        draw.line([(50, y), (760, y)], fill=(222, 226, 230), width=1)
    draw.ellipse((380, 260, 420, 300), outline=(196, 36, 36), width=6)
    draw.line([(400, 300), (400, 380)], fill=(196, 36, 36), width=6)
    draw.ellipse((390, 370, 410, 390), fill=(196, 36, 36))
    try:
        normal = ImageFont.truetype("arial.ttf", 28)
        bold = ImageFont.truetype("arialbd.ttf", 36)
    except OSError:
        normal = bold = ImageFont.load_default()
    draw.text((800, 100), "LOCALIZAÇÃO DA UC", fill="black", font=bold)
    draw.text((800, 180), f"Latitude: {latitude}", fill="black", font=normal)
    draw.text((800, 230), f"Longitude: {longitude}", fill="black", font=normal)
    draw.text((800, 310), "Abrir no Google Maps:", fill="black", font=normal)
    draw.text(
        (800, 360),
        f"google.com/maps?q={latitude},{longitude}",
        fill=(20, 75, 140),
        font=normal,
    )
    image.save(output)
    return output


def _insert_picture_at_placeholder(docx_path: Path, placeholder: str, image_path: Path, width: float = 6.2) -> bool:
    document = Document(docx_path)
    candidates = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                candidates.extend(cell.paragraphs)
    paragraph = next((item for item in candidates if placeholder in item.text), None)
    if paragraph is None:
        return False
    for run in paragraph.runs:
        run.text = run.text.replace(placeholder, "")
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    document.save(docx_path)
    return True


def _insert_location_figure(docx_path: Path, latitude: Any, longitude: Any) -> bool:
    with tempfile.TemporaryDirectory() as temp_dir:
        figure = _make_location_figure(latitude, longitude, Path(temp_dir) / "localizacao.png")
        return _insert_picture_at_placeholder(docx_path, "{{FIG_LOCALIZACAO}}", figure, width=6.2)


def _generation_series(workbook) -> tuple[list[float], list[float], float]:
    calculated = _calculated_values(workbook)
    consumption = [
        _number(workbook["Consumo_24m"][f"B{row}"].value) for row in range(2, 14)
    ]
    if not any(value is not None for value in consumption):
        consumption = [calculated.get(f"{{{{CONS_MES{idx}}}}}") for idx in range(1, 13)]
    consumption = [float(value or 0) for value in consumption]
    kwp = (
        _number(workbook["Inputs"]["B12"].value)
        or calculated.get("{{FV_POT_STC_KWP}}")
        or calculated.get("{{FV_POT_VI_KW}}")
        or 0
    )
    losses = _number(workbook["Inputs"]["B67"].value) if workbook["Inputs"].max_row >= 67 else None
    if losses is None:
        losses = 0.10
    generation = [float(kwp) * irr * 30 * (1 - losses) for irr in MONTHLY_IRRADIATION]
    return consumption, generation, float(kwp)


def _make_generation_chart(consumption: list[float], generation: list[float], kwp: float, output: Path) -> Path:
    width, height = 1450, 760
    margin_left, margin_right = 105, 50
    margin_top, margin_bottom = 95, 105
    plot_w = width - margin_left - margin_right
    plot_h = height - margin_top - margin_bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 42)
        font = ImageFont.truetype("arial.ttf", 24)
        small = ImageFont.truetype("arial.ttf", 21)
        bold = ImageFont.truetype("arialbd.ttf", 24)
    except OSError:
        title_font = font = small = bold = ImageFont.load_default()

    draw.text((margin_left, 28), "Estimativa de Consumo x Geração (kWh/mês)", fill=(30, 30, 30), font=title_font)
    draw.text((margin_left, 72), f"Sistema fotovoltaico: {kwp:.2f} kWp", fill=(80, 80, 80), font=small)

    max_value = max([1] + consumption + generation)
    step = max(100, int((max_value / 5 + 99) // 100) * 100)
    y_max = step * 5
    for i in range(6):
        value = step * i
        y = margin_top + plot_h - (value / y_max) * plot_h
        draw.line((margin_left, y, width - margin_right, y), fill=(225, 230, 235), width=2)
        draw.text((20, y - 12), f"{int(value)}", fill=(90, 90, 90), font=small)
    draw.line((margin_left, margin_top, margin_left, margin_top + plot_h), fill=(70, 70, 70), width=2)
    draw.line((margin_left, margin_top + plot_h, width - margin_right, margin_top + plot_h), fill=(70, 70, 70), width=2)

    group_w = plot_w / 12
    bar_w = group_w * 0.34
    blue = (47, 117, 181)
    orange = (237, 125, 49)
    for idx, month in enumerate(MONTHS):
        center = margin_left + group_w * idx + group_w / 2
        for value, color, offset in [
            (consumption[idx], blue, -bar_w * 0.55),
            (generation[idx], orange, bar_w * 0.55),
        ]:
            x0 = center + offset - bar_w / 2
            x1 = center + offset + bar_w / 2
            y0 = margin_top + plot_h - (value / y_max) * plot_h
            y1 = margin_top + plot_h
            draw.rectangle((x0, y0, x1, y1), fill=color)
        draw.text((center - 22, margin_top + plot_h + 18), month, fill=(50, 50, 50), font=small)

    legend_y = height - 48
    draw.rectangle((margin_left, legend_y, margin_left + 30, legend_y + 20), fill=blue)
    draw.text((margin_left + 42, legend_y - 3), "Consumo", fill=(50, 50, 50), font=bold)
    draw.rectangle((margin_left + 180, legend_y, margin_left + 210, legend_y + 20), fill=orange)
    draw.text((margin_left + 222, legend_y - 3), "Geração estimada", fill=(50, 50, 50), font=bold)
    draw.text((width - 360, legend_y - 3), "Fonte: tabela de dimensionamento", fill=(90, 90, 90), font=small)
    image.save(output)
    return output


def _insert_generation_chart(docx_path: Path, workbook) -> bool:
    with tempfile.TemporaryDirectory() as temp_dir:
        consumption, generation, kwp = _generation_series(workbook)
        chart = _make_generation_chart(consumption, generation, kwp, Path(temp_dir) / "grafico_geracao.png")
        return _insert_picture_at_placeholder(docx_path, "{{GRAFICO_GERACAO}}", chart, width=6.5)
    return True


def _unresolved_tags(docx_path: Path) -> list[str]:
    found: set[str] = set()
    with zipfile.ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                root = etree.fromstring(archive.read(name))
                text = "".join(node.text or "" for node in root.iter(W_T))
                found.update(TAG_RE.findall(text))
    return sorted(found)


def _document_tags(docx_path: Path) -> set[str]:
    return set(_unresolved_tags(docx_path))


def generate_memorial(excel: str | Path, template: str | Path, output: str | Path) -> GenerationResult:
    excel_path = Path(excel).resolve()
    template_path = Path(template).resolve()
    output_path = Path(output).resolve()
    if not excel_path.exists():
        raise GenerationError(f"Planilha não encontrada: {excel_path}")
    if not template_path.exists():
        raise GenerationError(f"Template não encontrado: {template_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)

    template_tags = _document_tags(template_path)
    replacements, missing = _read_values(excel_path)
    missing = [
        tag
        for tag in missing
        if tag in template_tags and not re.fullmatch(r"\{\{CARGA\d+_(?:DESC|P_W|QTD|FP|CI_KW|CI_KVA)\}\}", tag)
    ]
    workbook = load_workbook(excel_path, data_only=False)
    latitude = workbook["Inputs"]["B10"].value
    longitude = workbook["Inputs"]["B11"].value
    if latitude not in (None, "") and longitude not in (None, ""):
        inserted = _insert_location_figure(output_path, latitude, longitude)
        if not inserted:
            replacements["{{FIG_LOCALIZACAO}}"] = (
                f"Localização: https://www.google.com/maps?q={latitude},{longitude}"
            )
    else:
        replacements["{{FIG_LOCALIZACAO}}"] = ""
    if not _insert_generation_chart(output_path, workbook):
        replacements["{{GRAFICO_GERACAO}}"] = ""

    replaced = _replace_all_xml_parts(output_path, replacements)
    unresolved = _unresolved_tags(output_path)
    return GenerationResult(output_path, replaced, missing, unresolved)
