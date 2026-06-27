from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
ASSET_TEMPLATE = (
    ROOT
    / "outputs"
    / "GeradorMemorialSolar"
    / "assets"
    / "TEMPLATE_MEMORIAL_LIGHT_100_PARAM_TAGUEADO.docx"
)
SOURCE_TEMPLATE = (
    ROOT
    / "outputs"
    / "GeradorMemorialSolar"
    / "TEMPLATE_MEMORIAL_LIGHT_100_PARAM_TAGUEADO.docx"
)


def replace_in_runs(paragraph, old: str, new: str) -> None:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


document = Document(ASSET_TEMPLATE)

for paragraph in document.paragraphs:
    if paragraph.text.strip().startswith("Figura 1: Localização da unidade consumidora"):
        paragraph.text = (
            "Figura 1: Localização da unidade consumidora "
            "(imagem de satélite do Google Maps)."
        )

# Remove apenas o mapa fixo logo após {{FIG_LOCALIZACAO}}. A busca limitada
# torna esta atualização segura para ser executada mais de uma vez.
paragraphs = list(document.paragraphs)
marker_index = next(
    (
        index
        for index, paragraph in enumerate(paragraphs)
        if "{{FIG_LOCALIZACAO}}" in paragraph.text
    ),
    None,
)
if marker_index is not None:
    for paragraph in paragraphs[marker_index + 1 : marker_index + 5]:
        if paragraph._p.xpath(".//a:blip"):
            paragraph._p.getparent().remove(paragraph._p)
            break

# O gráfico nativo antigo é removido; o gerador insere apenas a imagem calculada
# no marcador {{GRAFICO_GERACAO}}.
for paragraph in list(document.paragraphs):
    if paragraph._p.xpath(".//*[local-name()='chart']"):
        paragraph._p.getparent().remove(paragraph._p)

inverter_table = document.tables[4]
inverter_table.cell(7, 0).text = "Máxima corrente CC – Icc-máx [A]"

document.save(ASSET_TEMPLATE)
document.save(SOURCE_TEMPLATE)
