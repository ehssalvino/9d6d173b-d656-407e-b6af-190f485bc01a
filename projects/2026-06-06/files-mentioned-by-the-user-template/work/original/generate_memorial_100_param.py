#!/usr/bin/env python3
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
import qrcode
from PIL import Image, ImageDraw, ImageFont

def read_map(ws):
    m = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or not row[0] or not row[1]:
            continue
        m[str(row[0]).strip()] = str(row[1]).strip()
    return m

def get(wb, ref):
    sheet, cell = ref.split('!')
    return wb[sheet][cell].value

def fmt(v):
    if v is None:
        return ''
    if isinstance(v, float):
        s = f"{v:.6f}".rstrip('0').rstrip('.')
        return s.replace('.', ',')
    return str(v)

def replace_all(doc, kv):
    for p in doc.paragraphs:
        if not p.runs:
            continue
        text = ''.join(r.text for r in p.runs)
        new = text
        for k,v in kv.items():
            if k in new:
                new = new.replace(k,v)
        if new != text:
            for r in p.runs: r.text=''
            p.runs[0].text = new
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.runs:
                        continue
                    text = ''.join(r.text for r in p.runs)
                    new = text
                    for k,v in kv.items():
                        if k in new:
                            new = new.replace(k,v)
                    if new != text:
                        for r in p.runs: r.text=''
                        p.runs[0].text = new

def make_fig(lat, lon, out_png):
    url = f"https://www.google.com/maps?q={lat},{lon}"
    qr = qrcode.QRCode(border=1, box_size=6)
    qr.add_data(url); qr.make(fit=True)
    qr_img = qr.make_image(fill_color='black', back_color='white').convert('RGB')
    W,H=1200,700
    img=Image.new('RGB',(W,H),'white')
    d=ImageDraw.Draw(img)
    for x in range(50,750,50): d.line([(x,50),(x,650)], fill=(220,220,220), width=1)
    for y in range(50,650,50): d.line([(50,y),(750,y)], fill=(220,220,220), width=1)
    d.ellipse((380,320,420,360), outline='red', width=5)
    d.line([(400,360),(400,430)], fill='red', width=5)
    d.ellipse((392,422,408,438), fill='red')
    try:
        font=ImageFont.truetype('DejaVuSans.ttf', 28)
        fontb=ImageFont.truetype('DejaVuSans-Bold.ttf', 34)
    except Exception:
        font=fontb=None
    d.text((800,80),'LOCALIZAÇÃO DA UC', fill='black', font=fontb)
    d.text((800,140), f'Lat: {lat}', fill='black', font=font)
    d.text((800,180), f'Lon: {lon}', fill='black', font=font)
    d.text((800,240), 'QR Code (Google Maps):', fill='black', font=font)
    img.paste(qr_img.resize((280,280)), (830,320))
    img.save(out_png)
    return out_png

def insert_fig(doc, png):
    ph='{{FIG_LOCALIZACAO}}'
    for p in doc.paragraphs:
        if ph in p.text:
            p.text = p.text.replace(ph,'')
            p.add_run().add_picture(png, width=Inches(6.2))
            return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if ph in p.text:
                        p.text = p.text.replace(ph,'')
                        p.add_run().add_picture(png, width=Inches(6.2))
                        return True
    return False

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument('--excel', required=True)
    ap.add_argument('--template', required=True)
    ap.add_argument('--out', required=True)
    args=ap.parse_args()

    wb = load_workbook(args.excel, data_only=True)
    mp = read_map(wb['Mapeamento'])

    kv = {}
    for ph, origin in mp.items():
        if ph == '{{FIG_LOCALIZACAO}}':
            continue
        kv[ph] = fmt(get(wb, origin))

    doc = Document(args.template)
    replace_all(doc, kv)

    lat = get(wb, 'Inputs!B10')
    lon = get(wb, 'Inputs!B11')
    if lat is not None and lon is not None:
        out_png = str(Path(args.out).with_suffix('.png'))
        make_fig(lat, lon, out_png)
        insert_fig(doc, out_png)
    else:
        replace_all(doc, {'{{FIG_LOCALIZACAO}}':''})

    doc.save(args.out)
    print('OK:', args.out)

if __name__=='__main__':
    main()
