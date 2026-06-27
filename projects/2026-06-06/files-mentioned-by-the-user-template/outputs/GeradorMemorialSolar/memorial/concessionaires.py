from __future__ import annotations

import math
import shutil
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from openpyxl import load_workbook

from .shared_data import expand_shared_values


LIGHT_FORMS = {
    "connection": "Formulario_de_Solicitacao_para_Orcamento_Conexao_MICRO_E_MINIGERACAO_.pdf",
    "generator": "Formulario_de_Micro-Minigeracao_-_Registro_Central_Geradora-2.pdf",
    "priority": "FROP_-_Formulario_de_Rateio_por_Ordem_de_Prioridade_SE_NECESSARIO_1.pdf",
    "percentage": "FRP_-_Formulario_de_Rateio_por_Percentuais_SE_NECESSARIO_1.pdf",
    "gd_data": "2022-01-31-dados-gd-ufv.docx",
}


@dataclass
class AllocationAccount:
    installation: str
    customer_code: str
    holder: str
    address: str
    percentage: str = ""


def _number(value: Any) -> float | None:
    if value in (None, ""):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace(" ", "")
    if "," in text and "." in text:
        text = text.replace(".", "").replace(",", ".")
    else:
        text = text.replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return None


def _text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _inputs(excel_path: str | Path) -> dict[str, Any]:
    workbook = load_workbook(excel_path, data_only=False)
    sheet = workbook["Inputs"]
    values = {
        _text(sheet[f"A{row}"].value): sheet[f"B{row}"].value
        for row in range(2, sheet.max_row + 1)
        if sheet[f"A{row}"].value
    }
    return expand_shared_values(values)


def _split_address(data: dict[str, Any]) -> dict[str, str]:
    street = _text(data.get("CLIENTE_ENDERECO_LOGRADOURO"))
    locality = _text(data.get("CLIENTE_BAIRRO_MUN_UF_CEP"))
    cep = ""
    import re

    match = re.search(r"\b\d{5}-?\d{3}\b", locality)
    if match:
        cep = match.group(0)
        locality = (locality[: match.start()] + locality[match.end() :]).strip(" ,-/")
    parts = [part.strip() for part in locality.split(",") if part.strip()]
    bairro = parts[0] if parts else ""
    city = ", ".join(parts[1:]) if len(parts) > 1 else _text(data.get("LOCAL_CIDADE_UF"))
    number = ""
    match = re.search(r"(?:,\s*|\s+)(\d+[A-Za-z]?)\s*$", street)
    if match:
        number = match.group(1)
        street = street[: match.start()].strip(" ,-")
    return {
        "street": street,
        "number": number,
        "bairro": bairro,
        "city": city,
        "cep": cep,
        "full": " ".join(
            item
            for item in (
                _text(data.get("CLIENTE_ENDERECO_LOGRADOURO")),
                _text(data.get("CLIENTE_BAIRRO_MUN_UF_CEP")),
            )
            if item
        ),
    }


def _dms(value: Any, negative_letter: str, positive_letter: str) -> str:
    number = _number(value)
    if number is None:
        return ""
    letter = negative_letter if number < 0 else positive_letter
    absolute = abs(number)
    degrees = int(absolute)
    total_seconds = round((absolute - degrees) * 3600, 2)
    minutes_value, seconds = divmod(total_seconds, 60)
    minutes = int(minutes_value)
    if minutes >= 60:
        degrees += 1
        minutes = 0
    return f"{degrees}° {minutes:02d}' {seconds:05.2f}\" {letter}"


def _area(data: dict[str, Any]) -> str:
    qty = _number(data.get("MOD_QTD"))
    dimensions = _text(data.get("MOD_DIMENSOES"))
    import re

    values = [float(item.replace(",", ".")) for item in re.findall(r"\d+(?:[.,]\d+)?", dimensions)]
    if qty and len(values) >= 2:
        width, height = values[0], values[1]
        if width > 20:
            width /= 1000
        if height > 20:
            height /= 1000
        return f"{qty * width * height:.2f}".replace(".", ",")
    return ""


def _font_size(value: str, width: float, default: float = 8) -> float:
    if not value:
        return default
    return max(5.2, min(default, width / max(len(value) * 0.52, 1)))


def _write(
    page: fitz.Page,
    rect: tuple[float, float, float, float],
    value: Any,
    *,
    size: float = 8,
    align: int = fitz.TEXT_ALIGN_LEFT,
    erase: bool = True,
) -> None:
    text = _text(value)
    if not text:
        return
    box = fitz.Rect(*rect)
    if erase:
        page.draw_rect(box, color=None, fill=(1, 1, 1), overlay=True)
    page.insert_textbox(
        box,
        text,
        fontsize=_font_size(text, box.width, size),
        fontname="helv",
        color=(0, 0, 0),
        align=align,
        overlay=True,
    )


def _check(page: fitz.Page, x: float, y: float, selected: bool) -> None:
    if not selected:
        return
    page.insert_text((x, y), "X", fontsize=9, fontname="helv", color=(0, 0, 0), overlay=True)


def _save_pdf(document: fitz.Document, output: Path) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output, garbage=4, deflate=True)
    document.close()
    return output


def _fill_connection_form(template: Path, output: Path, data: dict[str, Any], mode: str) -> Path:
    document = fitz.open(template)
    page = document[0]
    address = _split_address(data)
    tipo = _text(data.get("TIPO_LIGACAO")).upper()
    power = _text(data.get("FV_POT_KWP"))

    fields = [
        ((93, 127, 169, 143), data.get("UC_CODIGO_CLIENTE")),
        ((258, 127, 352, 143), data.get("UC_CONTA_CONTRATO")),
        ((548, 127, 570, 143), data.get("UC_CLASSE")),
        ((95, 152, 570, 168), data.get("CLIENTE_NOME")),
        ((78, 177, 460, 193), address["street"]),
        ((535, 177, 570, 193), address["number"]),
        ((52, 202, 123, 218), address["cep"]),
        ((188, 202, 311, 218), address["bairro"]),
        ((374, 202, 469, 218), address["city"]),
        ((94, 227, 290, 243), data.get("CLIENTE_CPF")),
        ((420, 227, 474, 243), data.get("CLIENTE_CELULAR")),
        ((530, 227, 570, 243), data.get("CLIENTE_TELEFONE")),
        ((62, 252, 570, 268), data.get("CLIENTE_EMAIL")),
        ((367, 278, 570, 294), f"{_dms(data.get('COORD_SUL'), 'S', 'N')}  {_dms(data.get('COORD_OESTE'), 'O', 'L')}"),
        ((147, 328, 203, 344), data.get("CARGA_INSTALADA_KW") or power),
        ((296, 328, 329, 344), data.get("INV_VAC_NOM") or data.get("DISJ_TENSAO_V")),
        ((99, 353, 143, 369), data.get("DISJ_CORRENTE_A")),
        ((240, 353, 275, 369), data.get("CONDUTOR_FASE_MM2")),
        ((265, 670, 326, 687), data.get("FV_POT_ATUAL_KW")),
        ((449, 670, 525, 687), power),
        ((272, 721, 390, 738), power),
    ]
    for rect, value in fields:
        _write(page, rect, value)

    installed_power = _number(data.get("FV_POT_KWP")) or 0
    _check(page, 129, 78, installed_power > 75)
    _check(page, 370, 78, installed_power <= 75)
    _check(page, 416, 338, "MONO" in tipo)
    _check(page, 475, 338, "BI" in tipo)
    _check(page, 535, 338, "TRI" in tipo)
    _check(page, 107, 531, True)
    _check(page, 90, 573, _text(data.get("TIPO_RAMAL")).upper() != "SUBTERRANEO")
    _check(page, 141, 573, _text(data.get("TIPO_RAMAL")).upper() == "SUBTERRANEO")
    request_type = _text(data.get("TIPO_SOLICITACAO")).lower()
    _check(page, 119, 597, "nova" in request_type)
    _check(page, 119, 614, not request_type or "sem" in request_type)
    _check(page, 119, 630, "com" in request_type and "alter" in request_type)
    _check(page, 194, 757, True)
    _check(page, 261, 773, True)

    page2 = document[1]
    _check(page2, 213, 71, mode == "Compensação local")
    _check(page2, 314, 71, mode == "Autoconsumo remoto")
    _check(page2, 421, 71, mode == "Múltiplas unidades consumidoras")
    _check(page2, 129, 88, mode == "Geração compartilhada")

    page3 = document[2]
    _write(page3, (133, 291, 570, 308), data.get("SOLICITANTE_NOME") or data.get("CLIENTE_NOME"))
    _write(page3, (65, 316, 335, 333), data.get("CLIENTE_EMAIL"))
    _write(page3, (473, 316, 570, 333), data.get("CLIENTE_CELULAR"))
    _write(page3, (59, 341, 218, 358), data.get("LOCAL_CIDADE_UF"))
    _write(page3, (260, 341, 343, 358), date.today().strftime("%d/%m/%Y"))
    _write(page3, (175, 392, 363, 409), data.get("RT_NOME"))
    _write(page3, (474, 392, 570, 409), data.get("RT_REGISTRO"))
    _write(page3, (75, 417, 194, 434), data.get("RT_ART_TRT"))
    _write(page3, (246, 417, 421, 434), data.get("RT_EMAIL"))
    _write(page3, (486, 417, 570, 434), data.get("RT_CELULAR"))
    _write(page3, (59, 442, 218, 459), data.get("LOCAL_CIDADE_UF"))
    _write(page3, (260, 442, 343, 459), date.today().strftime("%d/%m/%Y"))
    return _save_pdf(document, output)


def _fill_generator_form(template: Path, output: Path, data: dict[str, Any]) -> Path:
    document = fitz.open(template)
    page = document[0]
    values = [
        data.get("MOD_QTD"),
        data.get("MOD_MARCA"),
        data.get("MOD_MODELO"),
        _area(data),
        data.get("INV_QTD"),
        data.get("INV_MARCA"),
        data.get("INV_MODELO"),
        data.get("FV_POT_KWP"),
        data.get("INV_PN_KW"),
        f"{_dms(data.get('COORD_SUL'), 'S', 'N')} / {_dms(data.get('COORD_OESTE'), 'O', 'L')}",
        data.get("DATA_IMPLEMENTACAO"),
    ]
    rows = [
        (114, 130),
        (136, 152),
        (158, 174),
        (180, 196),
        (202, 218),
        (224, 240),
        (246, 262),
        (268, 284),
        (290, 306),
        (312, 334),
        (338, 356),
    ]
    for rect_y, value in zip(rows, values):
        _write(page, (276, rect_y[0], 408, rect_y[1]), value, size=7.4)
    return _save_pdf(document, output)


def _draw_allocation_table(
    page: fitz.Page,
    accounts: list[AllocationAccount],
    *,
    percentage: bool,
) -> None:
    top = 306
    left, right = 85, 513
    bottom = min(438, top + 18 * max(4, len(accounts)))
    page.draw_rect(fitz.Rect(left, top, right, bottom), color=None, fill=(1, 1, 1), overlay=True)
    widths = [92, 210, 82, 44]
    xs = [left]
    for width in widths:
        xs.append(xs[-1] + width)
    row_height = min(18, (bottom - top) / max(len(accounts), 1))
    for index, account in enumerate(accounts):
        y0 = top + index * row_height
        y1 = y0 + row_height
        for x in xs:
            page.draw_line((x, y0), (x, y1), color=(0, 0, 0), width=0.5, overlay=True)
        page.draw_line((left, y0), (right, y0), color=(0, 0, 0), width=0.5, overlay=True)
        values = [
            account.customer_code or account.installation,
            account.address,
            "Principal" if index == 0 else "Compensação",
            account.percentage if percentage else str(index),
        ]
        for column, value in enumerate(values):
            _write(
                page,
                (xs[column] + 2, y0 + 2, xs[column + 1] - 2, y1 - 1),
                value,
                size=7,
                erase=False,
                align=fitz.TEXT_ALIGN_CENTER if column in (2, 3) else fitz.TEXT_ALIGN_LEFT,
            )
    page.draw_line((left, bottom), (right, bottom), color=(0, 0, 0), width=0.5, overlay=True)


def _fill_priority_form(
    template: Path,
    output: Path,
    data: dict[str, Any],
    accounts: list[AllocationAccount],
    remainder_installation: str,
) -> Path:
    document = fitz.open(template)
    page = document[0]
    primary = AllocationAccount(
        installation=_text(data.get("UC_CONTA_CONTRATO")),
        customer_code=_text(data.get("UC_CODIGO_CLIENTE")),
        holder=_text(data.get("CLIENTE_NOME")),
        address=_split_address(data)["full"],
    )
    _draw_allocation_table(page, [primary, *accounts], percentage=False)
    _write(page, (225, 475, 455, 495), remainder_installation, size=9)
    locality = _text(data.get("LOCAL_CIDADE_UF"))
    date_text = date.today().strftime("%d/%m/%Y")
    _write(page, (92, 568, 505, 588), f"{locality + ', ' if locality else ''}{date_text}", size=9)
    return _save_pdf(document, output)


def _fill_percentage_form(
    template: Path,
    output: Path,
    data: dict[str, Any],
    accounts: list[AllocationAccount],
) -> Path:
    document = fitz.open(template)
    page = document[0]
    primary = AllocationAccount(
        installation=_text(data.get("UC_CONTA_CONTRATO")),
        customer_code=_text(data.get("UC_CODIGO_CLIENTE")),
        holder=_text(data.get("CLIENTE_NOME")),
        address=_split_address(data)["full"],
        percentage=_text(data.get("RATEIO_PRINCIPAL_PERCENTUAL")),
    )
    _draw_allocation_table(page, [primary, *accounts], percentage=True)
    locality = _text(data.get("LOCAL_CIDADE_UF"))
    date_text = date.today().strftime("%d/%m/%Y")
    _write(page, (92, 441, 505, 461), f"{locality + ', ' if locality else ''}{date_text}", size=9)
    return _save_pdf(document, output)


def _replace_paragraph_value(document: Document, label: str, value: Any) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(label):
            paragraph.text = f"{label} {_text(value)}"
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
            return


def _fill_gd_data_docx(
    template: Path,
    output: Path,
    data: dict[str, Any],
    mode: str,
    allocation_count: int,
) -> Path:
    shutil.copy2(template, output)
    document = Document(output)
    address = _split_address(data)
    values = {
        "Modalidade:": mode,
        "Quantidade de UCs que recebem os créditos:": allocation_count,
        "Classe:": data.get("UC_CLASSE"),
        "Município/UF da UC com GD:": data.get("LOCAL_CIDADE_UF"),
        "Endereço da UC com GD:": data.get("CLIENTE_ENDERECO_LOGRADOURO"),
        "CEP da UC com GD:": address["cep"],
        "Latitude:": _dms(data.get("COORD_SUL"), "S", "N"),
        "Longitude:": _dms(data.get("COORD_OESTE"), "O", "L"),
        "CPF/CNPJ do Titular:": data.get("CLIENTE_CPF"),
        "Nome do Titular da UC com GD:": data.get("CLIENTE_NOME"),
        "Telefone do Titular (DDD + número):": data.get("CLIENTE_CELULAR"),
        "E-mail do Titular:": data.get("CLIENTE_EMAIL"),
        "Município:": data.get("LOCAL_CIDADE_UF"),
        "Endereço:": address["full"],
        "CEP:": address["cep"],
        "Potência Total dos Módulos (kW):": data.get("FV_POT_KWP"),
        "Quantidade de Módulos:": data.get("MOD_QTD"),
        "Fabricante(s) dos Módulos:": data.get("MOD_MARCA"),
        "Modelo(s) dos Módulos:": data.get("MOD_MODELO"),
        "Potência Total dos Inversores (kW):": data.get("INV_PN_KW"),
        "Quantidade de Inversores:": data.get("INV_QTD"),
        "Fabricante(s) dos Inversores:": data.get("INV_MARCA"),
        "Modelo(s) dos Inversores:": data.get("INV_MODELO"),
        "Área Total dos Arranjos (m2):": _area(data),
        "Data da implantação da unidade geradora:": data.get("DATA_IMPLEMENTACAO"),
        "Data da conexão da unidade geradora na Distribuidora:": data.get("DATA_CONEXAO"),
        "CEG do empreendimento - GGG.FF.UF.999999-9.VV:": data.get("CEG_EMPREENDIMENTO"),
        "Nome da Usina:": data.get("USINA_NOME"),
        "Tipo do Ato de Outorga ou Registro:": data.get("OUTORGA_TIPO"),
        "Número do Ato de Outorga ou Registro:": data.get("OUTORGA_NUMERO"),
        "Ano do Ato de Outorga ou Registro:": data.get("OUTORGA_ANO"),
    }
    for label, value in values.items():
        _replace_paragraph_value(document, label, value)
    document.save(output)
    return output


def _make_enel_summary(
    output: Path,
    data: dict[str, Any],
    mode: str,
    accounts: list[AllocationAccount],
) -> Path:
    document = Document()
    document.add_heading("Pacote de dados para protocolo - Enel Distribuição Rio", level=1)
    document.add_paragraph(
        "Resumo de apoio ao preenchimento do portal/formulário vigente da concessionária. "
        "Este documento não substitui um formulário oficial da Enel."
    )
    sections = [
        (
            "Unidade consumidora",
            [
                ("Titular", data.get("CLIENTE_NOME")),
                ("CPF/CNPJ", data.get("CLIENTE_CPF")),
                ("Instalação/UC", data.get("UC_CONTA_CONTRATO")),
                ("Código do cliente", data.get("UC_CODIGO_CLIENTE")),
                ("Endereço", _split_address(data)["full"]),
                ("E-mail", data.get("CLIENTE_EMAIL")),
                ("Telefone", data.get("CLIENTE_CELULAR")),
            ],
        ),
        (
            "Central geradora",
            [
                ("Potência instalada", f"{_text(data.get('FV_POT_KWP'))} kW"),
                ("Módulos", f"{_text(data.get('MOD_QTD'))} x {_text(data.get('MOD_MARCA'))} {_text(data.get('MOD_MODELO'))}"),
                ("Inversores", f"{_text(data.get('INV_QTD'))} x {_text(data.get('INV_MARCA'))} {_text(data.get('INV_MODELO'))}"),
                ("Ligação", data.get("TIPO_LIGACAO")),
                ("Latitude", _dms(data.get("COORD_SUL"), "S", "N")),
                ("Longitude", _dms(data.get("COORD_OESTE"), "O", "L")),
                ("Modalidade", mode),
            ],
        ),
    ]
    for heading, rows in sections:
        document.add_heading(heading, level=2)
        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = _text(value)
    if accounts:
        document.add_heading("Rateio por ordem de prioridade", level=2)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Prioridade", "Instalação", "Código cliente", "Titular", "Endereço"]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
        primary = AllocationAccount(
            installation=_text(data.get("UC_CONTA_CONTRATO")),
            customer_code=_text(data.get("UC_CODIGO_CLIENTE")),
            holder=_text(data.get("CLIENTE_NOME")),
            address=_split_address(data)["full"],
        )
        for priority, account in enumerate([primary, *accounts]):
            cells = table.add_row().cells
            values = [priority, account.installation, account.customer_code, account.holder, account.address]
            for cell, value in zip(cells, values):
                cell.text = _text(value)
    document.save(output)
    return output


def generate_concessionaire_package(
    excel_path: str | Path,
    forms_root: str | Path,
    output_dir: str | Path,
    concessionaire: str,
    selected_forms: set[str],
    compensation_mode: str,
    accounts: list[AllocationAccount],
    remainder_installation: str = "",
) -> list[Path]:
    data = _inputs(excel_path)
    forms = Path(forms_root)
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []

    if concessionaire.upper().startswith("LIGHT"):
        if "connection" in selected_forms:
            generated.append(
                _fill_connection_form(
                    forms / LIGHT_FORMS["connection"],
                    output / "LIGHT_01_Solicitacao_Orcamento_Conexao.pdf",
                    data,
                    compensation_mode,
                )
            )
        if "generator" in selected_forms:
            generated.append(
                _fill_generator_form(
                    forms / LIGHT_FORMS["generator"],
                    output / "LIGHT_02_Registro_Central_Geradora.pdf",
                    data,
                )
            )
        if "gd_data" in selected_forms:
            generated.append(
                _fill_gd_data_docx(
                    forms / LIGHT_FORMS["gd_data"],
                    output / "LIGHT_03_Dados_GD_UFV.docx",
                    data,
                    compensation_mode,
                    len(accounts),
                )
            )
        if "priority" in selected_forms:
            generated.append(
                _fill_priority_form(
                    forms / LIGHT_FORMS["priority"],
                    output / "LIGHT_04_Rateio_Prioridade.pdf",
                    data,
                    accounts,
                    remainder_installation,
                )
            )
        if "percentage" in selected_forms:
            generated.append(
                _fill_percentage_form(
                    forms / LIGHT_FORMS["percentage"],
                    output / "LIGHT_05_Rateio_Percentuais.pdf",
                    data,
                    accounts,
                )
            )
    else:
        generated.append(
            _make_enel_summary(
                output / "ENEL_RJ_Pacote_Dados_Protocolo.docx",
                data,
                compensation_mode,
                accounts,
            )
        )
    return generated
