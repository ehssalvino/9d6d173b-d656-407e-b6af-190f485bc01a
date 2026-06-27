from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
SOURCE_DOCX = ASSETS / "TEMPLATE_MEMORIAL_LIGHT_100_PARAM_TAGUEADO.docx"
TARGET_DOCX = ASSETS / "TEMPLATE_MEMORIAL_ENEL_RJ_100_PARAM_TAGUEADO.docx"


PARAGRAPH_REPLACEMENTS = {
    89: (
        "O presente memorial técnico descritivo tem como objetivo apresentar à ENEL-RJ "
        "a metodologia e os dados técnicos para solicitação de conexão de microgeração "
        "distribuída fotovoltaica de {{FV_POT_KWP}} kWp, composta por {{INV_DESC}} "
        "{{TIPO_LIGACAO}} on-grid {{INV_MARCA}} modelo {{INV_MODELO}} (potência nominal "
        "{{INV_PN_KW}} kW) e {{MOD_QTD}} ({{MOD_QTD_TEXTO}}) módulos fotovoltaicos "
        "{{MOD_MARCA}} modelo {{MOD_MODELO}} ({{MOD_WP}} Wp cada), em instalação "
        "residencial individual, conforme a regulamentação da ANEEL, o PRODIST, as "
        "normas ABNT aplicáveis e as especificações técnicas vigentes da ENEL-RJ."
    ),
    98: (
        "ANEEL - Procedimentos de Distribuição de Energia Elétrica no Sistema Elétrico "
        "Nacional (PRODIST), especialmente o Módulo 3 - Conexão ao Sistema de Distribuição."
    ),
    99: (
        "Lei nº 14.300, de 6 de janeiro de 2022, que institui o marco legal da "
        "microgeração e minigeração distribuída."
    ),
    100: (
        "Resolução Normativa ANEEL nº 1.000, de 7 de dezembro de 2021, e suas alterações, "
        "incluindo as regras aplicáveis à conexão de microgeração distribuída."
    ),
    135: (
        "No ponto de entrega/conexão está instalado um disjuntor termomagnético, em "
        "conformidade com a Especificação Técnica 165 da ENEL-RJ e suas revisões vigentes, "
        "com as seguintes características:"
    ),
    155: (
        "NOTA 2: A potência de geração deve respeitar a potência disponibilizada para a "
        "unidade consumidora, conforme os critérios da Especificação Técnica 165 da ENEL-RJ."
    ),
    162: (
        "A caixa de medição EXISTENTE, POLIFÁSICA em material polimérico tem as dimensões "
        "de 380 mm x 470 mm x 190 mm (comprimento, altura e largura), está instalada no "
        "MURO, no ponto de entrega caracterizado como o limite da via pública com a "
        "propriedade, conforme fotos abaixo, atendendo aos requisitos de localização, "
        "acesso e lay-out da Especificação Técnica 165 da ENEL-RJ, conforme a FIGURA 2."
    ),
    213: (
        "Características da placa de advertência, conforme a Especificação Técnica 122 "
        "da ENEL-RJ e sua revisão vigente:"
    ),
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> None:
    dxf_source = Path(sys.argv[1]) if len(sys.argv) > 1 else None
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    document = Document(TARGET_DOCX)
    for index, text in PARAGRAPH_REPLACEMENTS.items():
        replace_paragraph_text(document.paragraphs[index], text)
    document.save(TARGET_DOCX)

    if dxf_source:
        target = ASSETS / "cad" / "TEMPLATE_DIAGRAMA_ENEL_RJ.dxf"
        shutil.copy2(dxf_source, target)
        print(target)
    print(TARGET_DOCX)


if __name__ == "__main__":
    main()
