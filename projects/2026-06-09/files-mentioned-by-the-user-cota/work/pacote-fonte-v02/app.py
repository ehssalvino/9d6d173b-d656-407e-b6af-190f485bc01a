from __future__ import annotations

import json
import os
import sqlite3
import subprocess
import sys
from dataclasses import asdict, dataclass, field
from datetime import date, datetime
from pathlib import Path
from tkinter import END, StringVar, Tk, filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


APP_NAME = "Delfos Propostas"
VERSION = "0.2.0"
BLUE = "31B5D0"
DARK = "17212B"
LIGHT = "EAF7FA"
GRAY = "66727D"


def app_data_dir() -> Path:
    base = Path(os.environ.get("LOCALAPPDATA", Path.home()))
    directory = base / "DelfosPropostas"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def resource_path(relative: str) -> Path:
    base = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    return base / relative


def money(value: float) -> str:
    text = f"{value:,.2f}"
    return "R$ " + text.replace(",", "#").replace(".", ",").replace("#", ".")


def number(value: str) -> float:
    normalized = value.strip().replace("R$", "").replace(" ", "")
    if "," in normalized:
        normalized = normalized.replace(".", "").replace(",", ".")
    try:
        return float(normalized or 0)
    except ValueError:
        return 0.0


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


@dataclass
class ProposalData:
    area: str
    proposal_number: str
    client: str
    document: str
    contact: str
    email: str
    address: str
    city: str
    issue_date: str
    validity_days: int
    responsible: str
    responsible_phone: str
    responsible_email: str
    technical: dict[str, str] = field(default_factory=dict)
    items: list[dict] = field(default_factory=list)
    service_value: float = 0
    discount: float = 0
    payment_terms: str = ""
    deadline: str = ""
    included: list[str] = field(default_factory=list)
    excluded: list[str] = field(default_factory=list)
    notes: str = ""

    @property
    def material_total(self) -> float:
        return sum(item["qty"] * item["unit"] for item in self.items)

    @property
    def total(self) -> float:
        return max(0, self.material_total + self.service_value - self.discount)


class Database:
    def __init__(self, path: Path | None = None):
        self.path = path or app_data_dir() / "delfos_propostas.db"
        self.connection = sqlite3.connect(self.path)
        self.connection.row_factory = sqlite3.Row
        self._create_schema()

    def _create_schema(self) -> None:
        self.connection.executescript(
            """
            PRAGMA foreign_keys = ON;
            CREATE TABLE IF NOT EXISTS clients (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL COLLATE NOCASE,
                document TEXT NOT NULL DEFAULT '',
                contact TEXT NOT NULL DEFAULT '',
                email TEXT NOT NULL DEFAULT '',
                address TEXT NOT NULL DEFAULT '',
                city TEXT NOT NULL DEFAULT '',
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_clients_name ON clients(name);
            CREATE INDEX IF NOT EXISTS idx_clients_document ON clients(document);

            CREATE TABLE IF NOT EXISTS proposals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                client_id INTEGER NOT NULL,
                proposal_number TEXT NOT NULL,
                area TEXT NOT NULL,
                issue_date TEXT NOT NULL,
                total REAL NOT NULL DEFAULT 0,
                payload TEXT NOT NULL,
                docx_path TEXT NOT NULL DEFAULT '',
                pdf_path TEXT NOT NULL DEFAULT '',
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY(client_id) REFERENCES clients(id) ON DELETE CASCADE
            );
            CREATE INDEX IF NOT EXISTS idx_proposals_client ON proposals(client_id);
            CREATE INDEX IF NOT EXISTS idx_proposals_number ON proposals(proposal_number);
            """
        )
        self.connection.commit()

    def _find_client(self, data: ProposalData):
        if data.document:
            row = self.connection.execute(
                "SELECT id FROM clients WHERE document = ? LIMIT 1", (data.document,)
            ).fetchone()
            if row:
                return row["id"]
        row = self.connection.execute(
            "SELECT id FROM clients WHERE name = ? COLLATE NOCASE LIMIT 1", (data.client,)
        ).fetchone()
        return row["id"] if row else None

    def save_proposal(
        self,
        data: ProposalData,
        proposal_id: int | None = None,
        docx_path: str = "",
        pdf_path: str = "",
    ) -> int:
        now = datetime.now().isoformat(timespec="seconds")
        client_id = self._find_client(data)
        client_values = (
            data.client,
            data.document,
            data.contact,
            data.email,
            data.address,
            data.city,
            now,
        )
        if client_id:
            self.connection.execute(
                """
                UPDATE clients
                SET name=?, document=?, contact=?, email=?, address=?, city=?, updated_at=?
                WHERE id=?
                """,
                client_values + (client_id,),
            )
        else:
            cursor = self.connection.execute(
                """
                INSERT INTO clients
                (name, document, contact, email, address, city, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                client_values[:-1] + (now, now),
            )
            client_id = cursor.lastrowid

        payload = json.dumps(asdict(data), ensure_ascii=False)
        if proposal_id:
            existing = self.connection.execute(
                "SELECT docx_path, pdf_path FROM proposals WHERE id=?", (proposal_id,)
            ).fetchone()
            if not existing:
                proposal_id = None
            else:
                docx_path = docx_path or existing["docx_path"]
                pdf_path = pdf_path or existing["pdf_path"]

        if proposal_id:
            self.connection.execute(
                """
                UPDATE proposals
                SET client_id=?, proposal_number=?, area=?, issue_date=?, total=?,
                    payload=?, docx_path=?, pdf_path=?, updated_at=?
                WHERE id=?
                """,
                (
                    client_id,
                    data.proposal_number,
                    data.area,
                    data.issue_date,
                    data.total,
                    payload,
                    docx_path,
                    pdf_path,
                    now,
                    proposal_id,
                ),
            )
        else:
            cursor = self.connection.execute(
                """
                INSERT INTO proposals
                (client_id, proposal_number, area, issue_date, total, payload,
                 docx_path, pdf_path, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    client_id,
                    data.proposal_number,
                    data.area,
                    data.issue_date,
                    data.total,
                    payload,
                    docx_path,
                    pdf_path,
                    now,
                    now,
                ),
            )
            proposal_id = cursor.lastrowid
        self.connection.commit()
        return int(proposal_id)

    def list_proposals(self, search: str = "") -> list[sqlite3.Row]:
        term = f"%{search.strip()}%"
        return self.connection.execute(
            """
            SELECT p.id, c.name AS client, c.contact, p.proposal_number, p.area,
                   p.issue_date, p.total, p.docx_path, p.pdf_path, p.updated_at
            FROM proposals p
            JOIN clients c ON c.id = p.client_id
            WHERE c.name LIKE ? COLLATE NOCASE
               OR c.document LIKE ?
               OR c.contact LIKE ?
               OR p.proposal_number LIKE ?
            ORDER BY p.id DESC
            """,
            (term, term, term, term),
        ).fetchall()

    def get_proposal(self, proposal_id: int) -> dict:
        row = self.connection.execute(
            "SELECT payload FROM proposals WHERE id=?", (proposal_id,)
        ).fetchone()
        if not row:
            raise ValueError("Proposta não encontrada no banco de dados.")
        return json.loads(row["payload"])

    def get_proposal_record(self, proposal_id: int) -> sqlite3.Row:
        row = self.connection.execute(
            "SELECT * FROM proposals WHERE id=?", (proposal_id,)
        ).fetchone()
        if not row:
            raise ValueError("Proposta não encontrada no banco de dados.")
        return row

    def next_proposal_number(self) -> str:
        prefix = date.today().strftime("%Y%m%d")
        row = self.connection.execute(
            "SELECT COUNT(*) AS total FROM proposals WHERE proposal_number LIKE ?",
            (f"{prefix}-%",),
        ).fetchone()
        return f"{prefix}-{int(row['total']) + 1:03d}"

    def delete_proposal(self, proposal_id: int) -> None:
        row = self.connection.execute(
            "SELECT client_id FROM proposals WHERE id=?", (proposal_id,)
        ).fetchone()
        if not row:
            return
        client_id = row["client_id"]
        self.connection.execute("DELETE FROM proposals WHERE id=?", (proposal_id,))
        remaining = self.connection.execute(
            "SELECT 1 FROM proposals WHERE client_id=? LIMIT 1", (client_id,)
        ).fetchone()
        if not remaining:
            self.connection.execute("DELETE FROM clients WHERE id=?", (client_id,))
        self.connection.commit()

    def close(self) -> None:
        self.connection.close()


AREA_CONFIG = {
    "Energia Solar": {
        "fields": [
            ("consumo_mensal", "Consumo médio mensal (kWh)", "4100"),
            ("tarifa", "Tarifa de energia (R$/kWh)", "0,93"),
            ("tipo_rede", "Tipo de rede", "Trifásica"),
            ("potencia_sistema", "Potência do sistema (kWp)", "37,18"),
            ("geracao_mensal", "Geração mensal estimada (kWh)", "4823,55"),
            ("modulos", "Módulos fotovoltaicos", "52 x 715 W"),
            ("inversores", "Inversores", "2 x 10 kW"),
        ],
        "included": [
            "Vistoria técnica e levantamento das condições de instalação",
            "Projeto elétrico e documentação técnica aplicável",
            "Fornecimento e instalação dos equipamentos descritos",
            "Configuração do monitoramento e comissionamento",
            "Orientação de uso e entrega técnica do sistema",
        ],
        "excluded": [
            "Reforço estrutural, alvenaria e adequações civis não descritas",
            "Alterações exigidas pela concessionária após vistoria",
            "Serviços ou equipamentos não relacionados nesta proposta",
        ],
    },
    "CFTV": {
        "fields": [
            ("cameras", "Quantidade de câmeras", "16"),
            ("resolucao", "Resolução", "Full HD 1080p"),
            ("gravador", "Gravador", "DVR 16 canais H.265"),
            ("armazenamento", "Armazenamento", "HD 2 TB"),
            ("infraestrutura", "Infraestrutura", "Cabeamento e caixas de passagem"),
        ],
        "included": [
            "Instalação, fixação e direcionamento das câmeras",
            "Lançamento e organização do cabeamento descrito",
            "Configuração do gravador, acesso local e remoto",
            "Testes, identificação dos pontos e orientação ao usuário",
        ],
        "excluded": [
            "Obras civis, pintura e recomposição de acabamento",
            "Internet, link de dados ou adequação da rede do cliente",
            "Pontos adicionais e materiais não descritos na relação de itens",
        ],
    },
    "Elétrica": {
        "fields": [
            ("servico", "Tipo de serviço", "Adequação de instalação elétrica"),
            ("tensao", "Tensão", "220 V"),
            ("fases", "Alimentação", "Trifásica"),
            ("carga", "Carga estimada", ""),
            ("quadro", "Quadro elétrico", "Fornecimento e montagem"),
        ],
        "included": [
            "Levantamento técnico e definição dos circuitos",
            "Fornecimento e instalação dos materiais descritos",
            "Identificação, testes elétricos e entrega funcional",
            "Limpeza básica da área de trabalho",
        ],
        "excluded": [
            "Obras civis e acabamentos não indicados no escopo",
            "Aumento de carga ou serviços da concessionária",
            "Correções em instalações existentes não previstas",
        ],
    },
    "Rede": {
        "fields": [
            ("pontos", "Quantidade de pontos", "24"),
            ("categoria", "Categoria do cabeamento", "CAT6"),
            ("rack", "Rack", "12U"),
            ("switches", "Switches", "1 x 24 portas Gigabit"),
            ("wifi", "Pontos de acesso Wi-Fi", "2"),
        ],
        "included": [
            "Lançamento, conectorização e identificação do cabeamento",
            "Organização do rack e instalação dos equipamentos descritos",
            "Certificação funcional dos pontos",
            "Configuração básica da rede e documentação dos pontos",
        ],
        "excluded": [
            "Serviço de provedor de internet e endereços públicos",
            "Obras civis, eletrocalhas ou infraestrutura não descrita",
            "Licenças de software e serviços recorrentes",
        ],
    },
}


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 27, DARK),
        ("Heading 1", 18, BLUE),
        ("Heading 2", 13, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(7)


def add_header_footer(doc: Document, data: ProposalData) -> None:
    for section in doc.sections:
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Inches(6.9))
        table.columns[0].width = Inches(4.8)
        table.columns[1].width = Inches(2.1)
        logo = resource_path("assets/logo.png")
        if logo.exists():
            run = table.cell(0, 0).paragraphs[0].add_run()
            run.add_picture(str(logo), width=Cm(3.2))
        right = table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = right.add_run(f"PROPOSTA {data.proposal_number}\n{data.area.upper()}")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Delfos Serviços | {data.responsible_phone} | {data.responsible_email}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)


def add_label_value_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11.5)
    for label, value in pairs:
        cells = table.add_row().cells
        cells[0].width = Cm(5)
        cells[1].width = Cm(11.5)
        set_cell_shading(cells[0], LIGHT)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_run = cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.color.rgb = RGBColor.from_string(DARK)
        cells[1].paragraphs[0].add_run(value or "-")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def generate_docx(data: ProposalData, destination: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_header_footer(doc, data)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(45)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("PROPOSTA COMERCIAL")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(data.area.upper())
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    client = doc.add_paragraph()
    client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client.paragraph_format.space_before = Pt(34)
    run = client.add_run(data.client)
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{data.city} | {data.issue_date}").font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()

    doc.add_heading("Apresentação", level=1)
    doc.add_paragraph(
        "A Delfos Serviços apresenta esta proposta técnica e comercial para atender "
        f"às necessidades de {data.area.lower()} informadas pelo cliente. O escopo "
        "foi organizado para facilitar a conferência dos serviços, equipamentos, "
        "condições comerciais e responsabilidades de cada parte."
    )

    doc.add_heading("Dados da proposta", level=2)
    add_label_value_table(
        doc,
        [
            ("Cliente", data.client),
            ("CPF/CNPJ", data.document),
            ("Contato", data.contact),
            ("E-mail", data.email),
            ("Local", ", ".join(filter(None, [data.address, data.city]))),
            ("Responsável Delfos", data.responsible),
            ("Validade", f"{data.validity_days} dias a partir da emissão"),
        ],
    )

    doc.add_heading("Resumo técnico", level=1)
    technical_pairs = []
    labels = {key: label for key, label, _ in AREA_CONFIG[data.area]["fields"]}
    for key, value in data.technical.items():
        technical_pairs.append((labels.get(key, key.replace("_", " ").title()), value))
    add_label_value_table(doc, technical_pairs)

    if data.area == "Energia Solar":
        consumption = number(data.technical.get("consumo_mensal", "0"))
        tariff = number(data.technical.get("tarifa", "0"))
        generation = number(data.technical.get("geracao_mensal", "0"))
        annual_savings = min(consumption, generation) * tariff * 12
        if annual_savings > 0:
            doc.add_heading("Estimativa de benefício", level=2)
            p = doc.add_paragraph()
            p.add_run("Economia anual estimada: ").bold = True
            p.add_run(money(annual_savings))
            doc.add_paragraph(
                "Estimativa simplificada baseada nos dados informados. O resultado "
                "real depende de irradiação, perdas, perfil de consumo, tarifas e "
                "regras da distribuidora."
            )

    doc.add_heading("Relação de equipamentos e materiais", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Qtd.", "Descrição", "Valor unitário", "Subtotal", "Observação"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top=120, bottom=120)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for item in data.items:
        cells = table.add_row().cells
        values = [
            f"{item['qty']:g}",
            item["description"],
            money(item["unit"]),
            money(item["qty"] * item["unit"]),
            item.get("note", ""),
        ]
        for index, value in enumerate(values):
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[index].paragraphs[0].add_run(value)

    doc.add_heading("Escopo dos serviços", level=1)
    doc.add_heading("Incluso", level=2)
    add_bullets(doc, data.included)
    doc.add_heading("Não incluso", level=2)
    add_bullets(doc, data.excluded)

    doc.add_heading("Investimento", level=1)
    add_label_value_table(
        doc,
        [
            ("Materiais e equipamentos", money(data.material_total)),
            ("Serviços", money(data.service_value)),
            ("Desconto", money(data.discount)),
            ("Valor total da proposta", money(data.total)),
        ],
    )

    doc.add_heading("Condições comerciais", level=2)
    add_label_value_table(
        doc,
        [
            ("Forma de pagamento", data.payment_terms),
            ("Prazo estimado", data.deadline),
            ("Validade da proposta", f"{data.validity_days} dias"),
            ("Garantia", "Conforme fabricante e legislação aplicável"),
        ],
    )

    if data.notes:
        doc.add_heading("Observações", level=2)
        doc.add_paragraph(data.notes)

    doc.add_heading("Aceite da proposta", level=1)
    doc.add_paragraph(
        "Ao assinar, o cliente declara que leu e concorda com o escopo, os valores, "
        "as condições comerciais e as exclusões apresentadas."
    )
    doc.add_paragraph("\n\n")
    sign = doc.add_table(rows=2, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign.cell(0, 0).paragraphs[0].add_run("_" * 42)
    sign.cell(0, 1).paragraphs[0].add_run("_" * 42)
    sign.cell(1, 0).paragraphs[0].add_run(data.client)
    sign.cell(1, 1).paragraphs[0].add_run("Delfos Serviços")
    for row in sign.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destination))


def convert_to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    escaped_docx = str(docx_path.resolve()).replace("'", "''")
    escaped_pdf = str(pdf_path.resolve()).replace("'", "''")
    command = (
        "$word=New-Object -ComObject Word.Application;"
        "$word.Visible=$false;"
        f"$doc=$word.Documents.Open('{escaped_docx}');"
        f"$doc.SaveAs([ref]'{escaped_pdf}',[ref]17);"
        "$doc.Close();$word.Quit();"
    )
    result = subprocess.run(
        ["powershell.exe", "-NoProfile", "-Command", command],
        capture_output=True,
        text=True,
        timeout=120,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )
    if result.returncode != 0 or not pdf_path.exists():
        raise RuntimeError(result.stderr.strip() or "Não foi possível gerar o PDF.")
    return pdf_path


class App:
    def __init__(self, root: Tk):
        self.root = root
        self.root.title(f"{APP_NAME} {VERSION}")
        self.root.geometry("1120x760")
        self.root.minsize(980, 680)
        self.db = Database()
        self.current_proposal_id: int | None = None
        self.area = StringVar(value="Energia Solar")
        self.history_search = StringVar()
        self.vars: dict[str, StringVar] = {}
        self.technical_vars: dict[str, StringVar] = {}
        self._build_style()
        self._build_ui()
        self._load_area()
        self.vars["proposal_number"].set(self.db.next_proposal_number())
        self.refresh_history()
        self.root.protocol("WM_DELETE_WINDOW", self.close)

    def _build_style(self) -> None:
        style = ttk.Style()
        try:
            style.theme_use("vista")
        except Exception:
            pass
        style.configure("Title.TLabel", font=("Segoe UI Semibold", 20), foreground="#17212B")
        style.configure("Sub.TLabel", font=("Segoe UI", 10), foreground="#66727D")
        style.configure("Accent.TButton", font=("Segoe UI Semibold", 10))
        style.configure("TLabelframe.Label", font=("Segoe UI Semibold", 10))

    def _build_ui(self) -> None:
        outer = ttk.Frame(self.root, padding=14)
        outer.pack(fill="both", expand=True)
        ttk.Label(outer, text="Gerador de Propostas Delfos", style="Title.TLabel").pack(anchor="w")
        ttk.Label(
            outer,
            text="Cadastre clientes, consulte o histórico e gere propostas em Word e PDF.",
            style="Sub.TLabel",
        ).pack(anchor="w", pady=(0, 10))

        top = ttk.Frame(outer)
        top.pack(fill="x", pady=(0, 8))
        ttk.Label(top, text="Área técnica:").pack(side="left")
        area_box = ttk.Combobox(
            top,
            textvariable=self.area,
            values=list(AREA_CONFIG),
            state="readonly",
            width=24,
        )
        area_box.pack(side="left", padx=8)
        area_box.bind("<<ComboboxSelected>>", lambda _event: self._load_area())
        ttk.Button(top, text="Nova proposta", command=self.new_proposal).pack(side="right")
        ttk.Button(top, text="Exportar backup", command=self.save_json).pack(side="right", padx=6)
        ttk.Button(top, text="Importar backup", command=self.load_json).pack(side="right")

        self.notebook = ttk.Notebook(outer)
        self.notebook.pack(fill="both", expand=True)
        self.general_tab = ttk.Frame(self.notebook, padding=12)
        self.technical_tab = ttk.Frame(self.notebook, padding=12)
        self.items_tab = ttk.Frame(self.notebook, padding=12)
        self.scope_tab = ttk.Frame(self.notebook, padding=12)
        self.history_tab = ttk.Frame(self.notebook, padding=12)
        self.notebook.add(self.general_tab, text="1. Cliente e proposta")
        self.notebook.add(self.technical_tab, text="2. Dados técnicos")
        self.notebook.add(self.items_tab, text="3. Itens e valores")
        self.notebook.add(self.scope_tab, text="4. Escopo e condições")
        self.notebook.add(self.history_tab, text="5. Clientes e histórico")

        self._build_general()
        self._build_items()
        self._build_scope()
        self._build_history()

        actions = ttk.Frame(outer)
        actions.pack(fill="x", pady=(10, 0))
        self.record_status = ttk.Label(actions, text="Nova proposta", style="Sub.TLabel")
        self.record_status.pack(side="left")
        ttk.Button(
            actions,
            text="Salvar no banco",
            style="Accent.TButton",
            command=self.save_to_database,
        ).pack(side="right", padx=(8, 0))
        ttk.Button(
            actions, text="Gerar DOCX", style="Accent.TButton", command=lambda: self.generate(False)
        ).pack(side="right")
        ttk.Button(
            actions, text="Gerar DOCX + PDF", style="Accent.TButton", command=lambda: self.generate(True)
        ).pack(side="right", padx=8)

    def _field(self, parent, row: int, col: int, key: str, label: str, default: str = ""):
        ttk.Label(parent, text=label).grid(row=row, column=col * 2, sticky="w", padx=(0, 6), pady=5)
        var = self.vars.setdefault(key, StringVar(value=default))
        entry = ttk.Entry(parent, textvariable=var, width=35)
        entry.grid(row=row, column=col * 2 + 1, sticky="ew", padx=(0, 18), pady=5)
        return entry

    def _build_general(self) -> None:
        for col in (1, 3):
            self.general_tab.columnconfigure(col, weight=1)
        today = date.today()
        self._field(self.general_tab, 0, 0, "proposal_number", "Número", today.strftime("%Y%m%d-001"))
        self._field(self.general_tab, 0, 1, "client", "Cliente")
        self._field(self.general_tab, 1, 0, "document", "CPF/CNPJ")
        self._field(self.general_tab, 1, 1, "contact", "Telefone/WhatsApp")
        self._field(self.general_tab, 2, 0, "email", "E-mail")
        self._field(self.general_tab, 2, 1, "address", "Endereço")
        self._field(self.general_tab, 3, 0, "city", "Cidade/UF")
        self._field(self.general_tab, 3, 1, "issue_date", "Data", today.strftime("%d/%m/%Y"))
        self._field(self.general_tab, 4, 0, "validity_days", "Validade (dias)", "5")
        self._field(self.general_tab, 5, 0, "responsible", "Responsável", "Paulo Arruda")
        self._field(self.general_tab, 5, 1, "responsible_phone", "Telefone Delfos")
        self._field(self.general_tab, 6, 0, "responsible_email", "E-mail Delfos")

    def _build_items(self) -> None:
        help_text = (
            "Digite um item por linha no formato: quantidade | descrição | valor unitário | observação\n"
            "Exemplo: 16 | Câmera Full HD 1080p | 109,00 | Uso externo"
        )
        ttk.Label(self.items_tab, text=help_text, style="Sub.TLabel").pack(anchor="w")
        self.items_text = ScrolledText(self.items_tab, height=17, font=("Consolas", 10))
        self.items_text.pack(fill="both", expand=True, pady=8)
        commercial = ttk.Frame(self.items_tab)
        commercial.pack(fill="x")
        for index, (key, label, default) in enumerate(
            [
                ("service_value", "Serviços (R$)", "0,00"),
                ("discount", "Desconto (R$)", "0,00"),
            ]
        ):
            ttk.Label(commercial, text=label).grid(row=0, column=index * 2, padx=(0, 5))
            var = self.vars.setdefault(key, StringVar(value=default))
            ttk.Entry(commercial, textvariable=var, width=16).grid(
                row=0, column=index * 2 + 1, padx=(0, 18)
            )

    def _build_scope(self) -> None:
        self.scope_tab.columnconfigure(0, weight=1)
        self.scope_tab.columnconfigure(1, weight=1)
        ttk.Label(self.scope_tab, text="Serviços inclusos (um por linha)").grid(
            row=0, column=0, sticky="w"
        )
        ttk.Label(self.scope_tab, text="Não inclusos (um por linha)").grid(
            row=0, column=1, sticky="w", padx=(12, 0)
        )
        self.included_text = ScrolledText(self.scope_tab, height=10)
        self.excluded_text = ScrolledText(self.scope_tab, height=10)
        self.included_text.grid(row=1, column=0, sticky="nsew", pady=5)
        self.excluded_text.grid(row=1, column=1, sticky="nsew", padx=(12, 0), pady=5)
        self.scope_tab.rowconfigure(1, weight=1)

        lower = ttk.Frame(self.scope_tab)
        lower.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(10, 0))
        lower.columnconfigure(1, weight=1)
        ttk.Label(lower, text="Forma de pagamento").grid(row=0, column=0, sticky="w", pady=4)
        self.vars.setdefault("payment_terms", StringVar(value="À vista ou conforme negociação"))
        ttk.Entry(lower, textvariable=self.vars["payment_terms"]).grid(
            row=0, column=1, sticky="ew", padx=8, pady=4
        )
        ttk.Label(lower, text="Prazo estimado").grid(row=1, column=0, sticky="w", pady=4)
        self.vars.setdefault("deadline", StringVar(value="A definir após o aceite"))
        ttk.Entry(lower, textvariable=self.vars["deadline"]).grid(
            row=1, column=1, sticky="ew", padx=8, pady=4
        )
        ttk.Label(lower, text="Observações").grid(row=2, column=0, sticky="nw", pady=4)
        self.notes_text = ScrolledText(lower, height=5)
        self.notes_text.grid(row=2, column=1, sticky="ew", padx=8, pady=4)

    def _build_history(self) -> None:
        search_bar = ttk.Frame(self.history_tab)
        search_bar.pack(fill="x", pady=(0, 8))
        ttk.Label(search_bar, text="Buscar cliente, documento, telefone ou proposta:").pack(
            side="left"
        )
        search_entry = ttk.Entry(search_bar, textvariable=self.history_search, width=42)
        search_entry.pack(side="left", padx=8)
        search_entry.bind("<KeyRelease>", lambda _event: self.refresh_history())
        ttk.Button(search_bar, text="Limpar", command=self.clear_history_search).pack(side="left")

        columns = ("client", "contact", "number", "area", "date", "total", "updated")
        self.history_tree = ttk.Treeview(
            self.history_tab, columns=columns, show="headings", selectmode="browse"
        )
        headings = {
            "client": "Cliente",
            "contact": "Contato",
            "number": "Proposta",
            "area": "Área",
            "date": "Data",
            "total": "Valor",
            "updated": "Atualizada em",
        }
        widths = {
            "client": 220,
            "contact": 120,
            "number": 120,
            "area": 120,
            "date": 90,
            "total": 105,
            "updated": 135,
        }
        for column in columns:
            self.history_tree.heading(column, text=headings[column])
            self.history_tree.column(column, width=widths[column], anchor="w")
        self.history_tree.column("total", anchor="e")
        self.history_tree.pack(fill="both", expand=True)
        self.history_tree.bind("<Double-1>", lambda _event: self.open_selected_proposal())

        buttons = ttk.Frame(self.history_tab)
        buttons.pack(fill="x", pady=(8, 0))
        ttk.Button(
            buttons, text="Abrir dados selecionados", command=self.open_selected_proposal
        ).pack(side="left")
        ttk.Button(buttons, text="Abrir DOCX", command=lambda: self.open_saved_file("docx")).pack(
            side="left", padx=6
        )
        ttk.Button(buttons, text="Abrir PDF", command=lambda: self.open_saved_file("pdf")).pack(
            side="left"
        )
        ttk.Button(buttons, text="Excluir proposta", command=self.delete_selected_proposal).pack(
            side="right"
        )
        ttk.Label(
            buttons, text=f"Banco local: {self.db.path}", style="Sub.TLabel"
        ).pack(side="right", padx=12)

    def _load_area(self) -> None:
        for child in self.technical_tab.winfo_children():
            child.destroy()
        self.technical_vars = {}
        self.technical_tab.columnconfigure(1, weight=1)
        config = AREA_CONFIG[self.area.get()]
        for row, (key, label, default) in enumerate(config["fields"]):
            ttk.Label(self.technical_tab, text=label).grid(
                row=row, column=0, sticky="w", padx=(0, 8), pady=6
            )
            var = StringVar(value=default)
            self.technical_vars[key] = var
            ttk.Entry(self.technical_tab, textvariable=var).grid(
                row=row, column=1, sticky="ew", pady=6
            )
        self.included_text.delete("1.0", END)
        self.included_text.insert("1.0", "\n".join(config["included"]))
        self.excluded_text.delete("1.0", END)
        self.excluded_text.insert("1.0", "\n".join(config["excluded"]))
        self._seed_items()

    def _seed_items(self) -> None:
        samples = {
            "Energia Solar": [
                "52 | Módulo fotovoltaico bifacial 715 W | 0,00 | Conforme cotação do fornecedor",
                "2 | Inversor monofásico 10 kW | 0,00 | Com monitoramento",
                "1 | Estrutura, proteções e cabeamento do sistema | 0,00 | Kit completo",
            ],
            "CFTV": [
                "1 | DVR 16 canais H.265 | 1.141,00 | ",
                "1 | HD Purple 2 TB | 549,90 | ",
                "16 | Câmera externa Full HD 1080p | 109,00 | ",
                "1 | Cabeamento, conectores e caixas de passagem | 1.500,00 | Estimativa",
            ],
            "Elétrica": [
                "1 | Quadro elétrico e dispositivos de proteção | 0,00 | ",
                "1 | Cabos, eletrodutos e acessórios | 0,00 | Conforme levantamento",
            ],
            "Rede": [
                "1 | Caixa de cabo CAT6 100% cobre | 0,00 | 305 m",
                "24 | Keystone RJ45 CAT6 | 0,00 | ",
                "1 | Rack 12U com acessórios | 0,00 | ",
                "1 | Switch 24 portas Gigabit | 0,00 | ",
            ],
        }
        self.items_text.delete("1.0", END)
        self.items_text.insert("1.0", "\n".join(samples[self.area.get()]))

    def _lines(self, widget: ScrolledText) -> list[str]:
        return [line.strip() for line in widget.get("1.0", END).splitlines() if line.strip()]

    def collect(self) -> ProposalData:
        items = []
        for index, line in enumerate(self._lines(self.items_text), start=1):
            parts = [part.strip() for part in line.split("|")]
            if len(parts) < 3:
                raise ValueError(f"Item {index}: use quantidade | descrição | valor unitário.")
            items.append(
                {
                    "qty": number(parts[0]),
                    "description": parts[1],
                    "unit": number(parts[2]),
                    "note": parts[3] if len(parts) > 3 else "",
                }
            )
        if not self.vars["client"].get().strip():
            raise ValueError("Informe o nome do cliente.")
        return ProposalData(
            area=self.area.get(),
            proposal_number=self.vars["proposal_number"].get().strip(),
            client=self.vars["client"].get().strip(),
            document=self.vars["document"].get().strip(),
            contact=self.vars["contact"].get().strip(),
            email=self.vars["email"].get().strip(),
            address=self.vars["address"].get().strip(),
            city=self.vars["city"].get().strip(),
            issue_date=self.vars["issue_date"].get().strip(),
            validity_days=int(number(self.vars["validity_days"].get()) or 5),
            responsible=self.vars["responsible"].get().strip(),
            responsible_phone=self.vars["responsible_phone"].get().strip(),
            responsible_email=self.vars["responsible_email"].get().strip(),
            technical={key: var.get().strip() for key, var in self.technical_vars.items()},
            items=items,
            service_value=number(self.vars["service_value"].get()),
            discount=number(self.vars["discount"].get()),
            payment_terms=self.vars["payment_terms"].get().strip(),
            deadline=self.vars["deadline"].get().strip(),
            included=self._lines(self.included_text),
            excluded=self._lines(self.excluded_text),
            notes=self.notes_text.get("1.0", END).strip(),
        )

    def apply_payload(self, payload: dict, proposal_id: int | None = None) -> None:
        self.area.set(payload.get("area", "Energia Solar"))
        self._load_area()
        for key, var in self.vars.items():
            if key in payload:
                var.set(str(payload[key]))
        for key, var in self.technical_vars.items():
            if key in payload.get("technical", {}):
                var.set(str(payload["technical"][key]))
        self.items_text.delete("1.0", END)
        self.items_text.insert(
            "1.0",
            "\n".join(
                f"{item['qty']} | {item['description']} | {item['unit']} | {item.get('note', '')}"
                for item in payload.get("items", [])
            ),
        )
        self.included_text.delete("1.0", END)
        self.included_text.insert("1.0", "\n".join(payload.get("included", [])))
        self.excluded_text.delete("1.0", END)
        self.excluded_text.insert("1.0", "\n".join(payload.get("excluded", [])))
        self.notes_text.delete("1.0", END)
        self.notes_text.insert("1.0", payload.get("notes", ""))
        self.current_proposal_id = proposal_id
        if proposal_id:
            self.record_status.config(
                text=f"Editando proposta #{proposal_id}: {payload.get('client', '')}"
            )
        else:
            self.record_status.config(text="Nova proposta")

    def save_to_database(self, show_message: bool = True) -> int | None:
        try:
            data = self.collect()
            self.current_proposal_id = self.db.save_proposal(
                data, proposal_id=self.current_proposal_id
            )
            self.record_status.config(
                text=f"Proposta #{self.current_proposal_id} salva: {data.client}"
            )
            self.refresh_history()
            if show_message:
                messagebox.showinfo(
                    APP_NAME,
                    f"Cliente e proposta salvos no banco local.\n\n"
                    f"Registro: #{self.current_proposal_id}",
                )
            return self.current_proposal_id
        except Exception as exc:
            if show_message:
                messagebox.showerror(APP_NAME, str(exc))
            return None

    def refresh_history(self) -> None:
        if not hasattr(self, "history_tree"):
            return
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)
        for row in self.db.list_proposals(self.history_search.get()):
            updated = row["updated_at"].replace("T", " ")[:16]
            self.history_tree.insert(
                "",
                END,
                iid=str(row["id"]),
                values=(
                    row["client"],
                    row["contact"],
                    row["proposal_number"],
                    row["area"],
                    row["issue_date"],
                    money(row["total"]),
                    updated,
                ),
            )

    def clear_history_search(self) -> None:
        self.history_search.set("")
        self.refresh_history()

    def selected_proposal_id(self) -> int:
        selection = self.history_tree.selection()
        if not selection:
            raise ValueError("Selecione uma proposta no histórico.")
        return int(selection[0])

    def open_selected_proposal(self) -> None:
        try:
            proposal_id = self.selected_proposal_id()
            self.apply_payload(self.db.get_proposal(proposal_id), proposal_id)
            self.notebook.select(self.general_tab)
        except Exception as exc:
            messagebox.showerror(APP_NAME, str(exc))

    def open_saved_file(self, kind: str) -> None:
        try:
            proposal_id = self.selected_proposal_id()
            record = self.db.get_proposal_record(proposal_id)
            path_value = record[f"{kind}_path"]
            if not path_value or not Path(path_value).exists():
                raise ValueError(
                    f"O arquivo {kind.upper()} não foi localizado. "
                    "Abra os dados e gere novamente a proposta."
                )
            os.startfile(path_value)
        except Exception as exc:
            messagebox.showerror(APP_NAME, str(exc))

    def delete_selected_proposal(self) -> None:
        try:
            proposal_id = self.selected_proposal_id()
            if not messagebox.askyesno(
                APP_NAME,
                "Excluir esta proposta do histórico?\n\n"
                "Os arquivos DOCX e PDF não serão apagados.",
            ):
                return
            self.db.delete_proposal(proposal_id)
            if self.current_proposal_id == proposal_id:
                self.current_proposal_id = None
                self.record_status.config(text="Nova proposta")
            self.refresh_history()
        except Exception as exc:
            messagebox.showerror(APP_NAME, str(exc))

    def new_proposal(self) -> None:
        self.current_proposal_id = None
        for key in ("client", "document", "contact", "email", "address", "city"):
            self.vars[key].set("")
        self.vars["proposal_number"].set(self.db.next_proposal_number())
        self.vars["issue_date"].set(date.today().strftime("%d/%m/%Y"))
        self.vars["validity_days"].set("5")
        self.vars["service_value"].set("0,00")
        self.vars["discount"].set("0,00")
        self.vars["payment_terms"].set("À vista ou conforme negociação")
        self.vars["deadline"].set("A definir após o aceite")
        self.notes_text.delete("1.0", END)
        self._load_area()
        self.record_status.config(text="Nova proposta")
        self.notebook.select(self.general_tab)

    def generate(self, with_pdf: bool) -> None:
        try:
            data = self.collect()
            default_name = (
                f"Proposta_{data.area.replace(' ', '_')}_{data.client.replace(' ', '_')}.docx"
            )
            selected = filedialog.asksaveasfilename(
                title="Salvar proposta",
                defaultextension=".docx",
                initialfile=default_name,
                filetypes=[("Documento Word", "*.docx")],
            )
            if not selected:
                return
            destination = Path(selected)
            generate_docx(data, destination)
            generated = [str(destination)]
            pdf_path = ""
            if with_pdf:
                pdf_path = str(convert_to_pdf(destination))
                generated.append(pdf_path)
            self.current_proposal_id = self.db.save_proposal(
                data,
                proposal_id=self.current_proposal_id,
                docx_path=str(destination.resolve()),
                pdf_path=pdf_path,
            )
            self.record_status.config(
                text=f"Proposta #{self.current_proposal_id} salva: {data.client}"
            )
            self.refresh_history()
            messagebox.showinfo(APP_NAME, "Proposta gerada:\n\n" + "\n".join(generated))
            os.startfile(destination.parent)
        except Exception as exc:
            messagebox.showerror(APP_NAME, str(exc))

    def save_json(self) -> None:
        try:
            data = self.collect()
            selected = filedialog.asksaveasfilename(
                title="Salvar dados da proposta",
                defaultextension=".json",
                filetypes=[("Dados da proposta", "*.json")],
            )
            if not selected:
                return
            Path(selected).write_text(
                json.dumps(data.__dict__, ensure_ascii=False, indent=2), encoding="utf-8"
            )
        except Exception as exc:
            messagebox.showerror(APP_NAME, str(exc))

    def load_json(self) -> None:
        selected = filedialog.askopenfilename(
            title="Abrir dados da proposta", filetypes=[("Dados da proposta", "*.json")]
        )
        if not selected:
            return
        try:
            payload = json.loads(Path(selected).read_text(encoding="utf-8"))
            self.apply_payload(payload)
        except Exception as exc:
            messagebox.showerror(APP_NAME, f"Não foi possível abrir o arquivo:\n{exc}")

    def close(self) -> None:
        self.db.close()
        self.root.destroy()


def main() -> None:
    root = Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
