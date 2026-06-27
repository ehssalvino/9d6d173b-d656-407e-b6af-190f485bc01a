from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
SOURCES = ROOT / "fontes"
OUT = ROOT / "extracted"
OUT.mkdir(parents=True, exist_ok=True)


def clean_text(value: str | None) -> str:
    return re.sub(r"[ \t]+", " ", value or "").strip()


def extract_pdf(path: Path) -> dict:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text() or "")
        pages.append({"page": index, "text": text})
    return {"file": path.name, "type": "pdf", "pages": pages}


def extract_docx(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = clean_text(paragraph.text)
        if text:
            paragraphs.append(
                {
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": text,
                }
            )

    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([clean_text(cell.text) for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
            }
        )

    return {
        "file": path.name,
        "type": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": sections,
        "inline_shapes": len(doc.inline_shapes),
    }


def shared_strings(archive: zipfile.ZipFile) -> list[str]:
    name = "xl/sharedStrings.xml"
    if name not in archive.namelist():
        return []
    root = ET.fromstring(archive.read(name))
    ns = {"s": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    values = []
    for item in root.findall("s:si", ns):
        text = "".join(node.text or "" for node in item.findall(".//s:t", ns))
        values.append(clean_text(text))
    return values


def extract_xlsx(path: Path) -> dict:
    ns = {
        "m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "p": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    with zipfile.ZipFile(path) as archive:
        strings = shared_strings(archive)
        workbook = ET.fromstring(archive.read("xl/workbook.xml"))
        rels = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        relationship_targets = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels.findall("p:Relationship", ns)
        }
        sheets = []
        for sheet in workbook.findall("m:sheets/m:sheet", ns):
            name = sheet.attrib["name"]
            rel_id = sheet.attrib[f"{{{ns['r']}}}id"]
            target = relationship_targets[rel_id].lstrip("/")
            if not target.startswith("xl/"):
                target = f"xl/{target}"
            xml = ET.fromstring(archive.read(target))
            cells = []
            for cell in xml.findall(".//m:c", ns):
                ref = cell.attrib.get("r", "")
                cell_type = cell.attrib.get("t", "")
                formula_node = cell.find("m:f", ns)
                value_node = cell.find("m:v", ns)
                inline_node = cell.find("m:is/m:t", ns)
                value = ""
                if inline_node is not None:
                    value = inline_node.text or ""
                elif value_node is not None:
                    value = value_node.text or ""
                    if cell_type == "s" and value.isdigit():
                        index = int(value)
                        value = strings[index] if index < len(strings) else value
                formula = formula_node.text if formula_node is not None else ""
                if value or formula:
                    cells.append(
                        {
                            "ref": ref,
                            "value": clean_text(value),
                            "formula": clean_text(formula),
                        }
                    )
            sheets.append({"name": name, "cells": cells})
        return {"file": path.name, "type": "xlsx", "sheets": sheets}


def main() -> None:
    results = []
    for path in sorted(SOURCES.iterdir()):
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            result = extract_pdf(path)
        elif suffix == ".docx":
            result = extract_docx(path)
        elif suffix == ".xlsx":
            result = extract_xlsx(path)
        else:
            continue
        results.append(result)
        output_path = OUT / f"{path.stem}.json"
        output_path.write_text(
            json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    summary = {
        "files": [
            {
                "file": result["file"],
                "type": result["type"],
                "pages": len(result.get("pages", [])),
                "paragraphs": len(result.get("paragraphs", [])),
                "tables": len(result.get("tables", [])),
                "sheets": [
                    {"name": sheet["name"], "populated_cells": len(sheet["cells"])}
                    for sheet in result.get("sheets", [])
                ],
            }
            for result in results
        ]
    }
    (OUT / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
